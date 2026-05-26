from copy import deepcopy
import io
import logging
import re

from docx import Document
from docx.oxml.ns import qn

from app.services.bookmark_manager import BookmarkManager

logger = logging.getLogger(__name__)


class ResumeInjector:
    """Handles injecting extracted points into resume templates with bookmarks."""

    def __init__(self):
        self.bookmark_manager = BookmarkManager()

    def extract_points_by_heading(self, processed_text):
        """
        Extract points from processed text organized by cycle.
        Works with Cycle format.
        """
        points_by_cycle = {}
        current_cycle = None
        all_points = []

        for raw_line in processed_text.split("\n"):
            line = raw_line.strip()
            if not line:
                continue

            cycle_match = re.match(r"Cycle\s+(\d+):", line, re.IGNORECASE)
            if cycle_match:
                current_cycle = int(cycle_match.group(1))
                points_by_cycle.setdefault(current_cycle, [])
                continue

            point_text = None

            if line.startswith(("\u2022", "\u00e2\u20ac\u00a2")):
                point_text = re.sub(r"^(?:\u2022|\u00e2\u20ac\u00a2)\s*", "", line).strip()
            elif line.startswith("-") and not line.startswith("--"):
                point_text = re.sub(r"^-\s*", "", line).strip()
            elif line.startswith("*") and not line.startswith("**"):
                point_text = re.sub(r"^\*\s*", "", line).strip()
            elif line.startswith("+"):
                point_text = re.sub(r"^\+\s*", "", line).strip()
            elif re.match(r"^\d+\.", line):
                point_text = re.sub(r"^\d+\.\s*", "", line).strip()
            elif current_cycle is not None and len(line) >= 20 and not line.startswith(("=", "_", "-", "Cycle")):
                point_text = line

            if point_text:
                all_points.append(point_text)
                target_cycle = current_cycle or 1
                points_by_cycle.setdefault(target_cycle, []).append(point_text)

        logger.debug(f"Extracted cycles: {list(points_by_cycle.keys())}")
        logger.debug(f"Total points: {len(all_points)}")

        return points_by_cycle, all_points

    def find_bookmark_paragraph(self, doc, bookmark_name):
        """Find the paragraph that contains a bookmark."""
        for para in doc.paragraphs:
            if bookmark_name in para._element.xml:
                return para

        for element in doc.element.iter():
            if bookmark_name in element.tag or (
                hasattr(element, "attrib")
                and any(bookmark_name in str(v) for v in element.attrib.values())
            ):
                parent = element.getparent()
                while parent is not None:
                    if "p" in parent.tag:
                        for para in doc.paragraphs:
                            if para._element == parent:
                                return para
                    parent = parent.getparent()

        return None

    def _find_insertion_paragraph(self, doc, bookmark_name):
        bookmark_para = self.find_bookmark_paragraph(doc, bookmark_name)

        if bookmark_para:
            return bookmark_para

        logger.warning(f"Bookmark paragraph not found for '{bookmark_name}', searching by pattern...")
        company_section_name = bookmark_name.replace("_Responsibilities", "")
        for para_idx, para in enumerate(doc.paragraphs):
            if "Responsibilities" in para.text and company_section_name in para.text:
                logger.info(f"Found matching paragraph by pattern at index {para_idx}")
                return doc.paragraphs[para_idx + 1] if para_idx + 1 < len(doc.paragraphs) else para

        for para_idx, para in enumerate(doc.paragraphs):
            if "Responsibilities" in para.text:
                logger.info(f"Found Responsibilities section at index {para_idx}")
                return doc.paragraphs[min(para_idx + 2, len(doc.paragraphs) - 1)]

        return None

    def copy_list_formatting(self, source_para, target_para):
        """Copy list/bullet formatting from source paragraph to target paragraph."""
        try:
            source_pPr = source_para._element.get_or_add_pPr()
            target_pPr = target_para._element.get_or_add_pPr()
            source_numPr = source_pPr.find(qn("w:numPr"))

            if source_numPr is not None:
                try:
                    target_numPr = target_pPr.find(qn("w:numPr"))
                    if target_numPr is not None:
                        target_pPr.remove(target_numPr)

                    target_pPr.append(deepcopy(source_numPr))
                except Exception:
                    self._apply_fallback_bullet_formatting(target_para)
        except Exception:
            try:
                self._apply_fallback_bullet_formatting(target_para)
            except Exception:
                pass

    def _apply_fallback_bullet_formatting(self, para):
        """Apply fallback bullet formatting when XML manipulation fails."""
        try:
            if para.runs:
                first_run = para.runs[0]
                if first_run.text and not first_run.text.startswith("\u2022"):
                    first_run.text = "\u2022 " + first_run.text
        except Exception:
            pass

    def _insert_points_after_paragraph(self, doc, reference_para, cycle_points):
        """Insert points after a reference paragraph while copying its formatting."""
        parent = reference_para._element.getparent()

        for offset, point_text in enumerate(cycle_points):
            ref_style_name = reference_para.style.name if reference_para.style else "Normal"
            new_para = doc.add_paragraph(point_text, style=ref_style_name)

            ref_pformat = reference_para.paragraph_format
            new_pformat = new_para.paragraph_format
            new_pformat.left_indent = ref_pformat.left_indent
            new_pformat.first_line_indent = ref_pformat.first_line_indent
            new_pformat.space_before = ref_pformat.space_before
            new_pformat.space_after = ref_pformat.space_after
            if ref_pformat.line_spacing:
                new_pformat.line_spacing = ref_pformat.line_spacing

            if reference_para.runs:
                ref_run = reference_para.runs[0]
                for run in new_para.runs:
                    if ref_run.font.name:
                        run.font.name = ref_run.font.name
                    if ref_run.font.size:
                        run.font.size = ref_run.font.size
                    if ref_run.font.bold is not None:
                        run.font.bold = ref_run.font.bold
                    if ref_run.font.italic is not None:
                        run.font.italic = ref_run.font.italic
                    if ref_run.font.color.rgb:
                        run.font.color.rgb = ref_run.font.color.rgb

            self.copy_list_formatting(reference_para, new_para)
            new_parent = new_para._element.getparent()
            new_parent.remove(new_para._element)

            try:
                ref_index = list(parent).index(reference_para._element)
                parent.insert(ref_index + 1 + offset, new_para._element)
            except (ValueError, IndexError):
                parent.append(new_para._element)

    def _clear_paragraph_text(self, para):
        """Clear visible text while preserving paragraph structure and bookmarks."""
        for run in para.runs:
            run.text = ""

    def inject_points_into_resume(self, resume_bytes, processed_text, custom_mapping=None, unused_handling="keep"):
        """Inject extracted points into resume at bookmarks with flexible mapping."""
        try:
            # Ensure document has bookmarks (recover from reference if needed)
            resume_bytes, available_bookmarks, bookmark_metadata = self.bookmark_manager.ensure_bookmarks_from_reference(resume_bytes)
            if bookmark_metadata.get("auto_created"):
                logger.info(f"Auto-recovered bookmarks from reference: {bookmark_metadata}")

            try:
                if isinstance(resume_bytes, (str, bytes)):
                    if isinstance(resume_bytes, str) and not resume_bytes.startswith(
                        b"\x50\x4b\x03\x04".decode("latin1", errors="ignore")
                    ):
                        doc = Document(resume_bytes)
                    else:
                        doc = Document(io.BytesIO(resume_bytes) if isinstance(resume_bytes, bytes) else resume_bytes)
                else:
                    doc = Document(resume_bytes)
            except Exception as e:
                logger.error(f"Failed to load DOCX document: {str(e)}")
                raise ValueError(f"Invalid or corrupted DOCX file: {str(e)}.")

            if hasattr(resume_bytes, "seek"):
                resume_bytes.seek(0)

            points_by_cycle, all_points = self.extract_points_by_heading(processed_text)
            logger.info(f"Extracted points - Cycles: {len(points_by_cycle)}, Total points: {len(all_points)}")

            if not all_points:
                logger.error("No points found in processed text. Check the format.")
                raise ValueError("No points found in processed text. Check the format.")

            non_empty_cycles = {c: p for c, p in points_by_cycle.items() if p}
            if not non_empty_cycles:
                logger.error("No actual points found in any cycles. Cycles are empty.")
                raise ValueError("No actual points found in any cycles. Cycles are empty.")

            logger.info(f"Available bookmarks in resume: {available_bookmarks}")
            if not available_bookmarks:
                logger.error("No bookmarks found in resume template and recovery failed. Please add bookmarks first.")
                raise ValueError("No bookmarks found in resume template and recovery failed. Please add bookmarks first.")

            if unused_handling not in {"keep", "repeat", "clear"}:
                raise ValueError("unused_handling must be one of: keep, repeat, clear")

            if custom_mapping:
                cycle_to_bookmark = {int(k): v for k, v in custom_mapping.items() if v}
                is_valid, error_msg = self.bookmark_manager.validate_mapping(
                    cycle_to_bookmark, available_bookmarks
                )
                if not is_valid:
                    logger.error(f"Invalid mapping: {error_msg}")
                    raise ValueError(f"Invalid mapping: {error_msg}")
            else:
                num_cycles = len(points_by_cycle)
                cycle_to_bookmark = self.bookmark_manager.suggest_mappings(
                    available_bookmarks, num_cycles
                )
                logger.info(f"Suggested mapping: {cycle_to_bookmark}")

            if not cycle_to_bookmark:
                logger.error("Could not generate bookmark mappings. Please provide custom mapping.")
                raise ValueError("Could not generate bookmark mappings. Please provide custom mapping.")

            injections = {}
            injected_bookmarks = set()

            for cycle_num in sorted(points_by_cycle.keys()):
                if cycle_num not in cycle_to_bookmark:
                    logger.warning(f"No mapping for cycle {cycle_num}")
                    continue

                bookmark_name = cycle_to_bookmark[cycle_num]
                cycle_points = points_by_cycle[cycle_num]
                logger.info(
                    f"Injecting {len(cycle_points)} points for cycle {cycle_num} "
                    f"into bookmark '{bookmark_name}'"
                )

                if not cycle_points:
                    continue

                bookmark_para = self._find_insertion_paragraph(doc, bookmark_name)
                if bookmark_para:
                    self._insert_points_after_paragraph(doc, bookmark_para, cycle_points)
                    injections[bookmark_name] = len(cycle_points)
                    injected_bookmarks.add(bookmark_name)

            unused_bookmarks = [bm for bm in available_bookmarks if bm not in injected_bookmarks]
            if unused_handling == "repeat" and unused_bookmarks:
                last_cycle_points = non_empty_cycles[max(non_empty_cycles.keys())]
                for bookmark_name in unused_bookmarks:
                    bookmark_para = self._find_insertion_paragraph(doc, bookmark_name)
                    if bookmark_para:
                        self._insert_points_after_paragraph(doc, bookmark_para, last_cycle_points)
                        injections[bookmark_name] = len(last_cycle_points)
            elif unused_handling == "clear" and unused_bookmarks:
                for bookmark_name in unused_bookmarks:
                    bookmark_para = self.find_bookmark_paragraph(doc, bookmark_name)
                    if bookmark_para:
                        self._clear_paragraph_text(bookmark_para)
                        injections[bookmark_name] = 0

            if not injections:
                logger.error("Failed to inject points. No valid insertion points found.")
                raise ValueError("Failed to inject points. No valid insertion points found.")

            logger.info(f"Successfully injected points into {len(injections)} bookmarks")
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)

            return output, injections

        except Exception as e:
            logger.error(f"Error injecting points into resume: {str(e)}", exc_info=True)
            raise Exception(f"Error injecting points into resume: {str(e)}")

    def get_available_bookmarks(self, resume_bytes):
        """Get list of available bookmarks in resume template."""
        return self.bookmark_manager.detect_bookmarks(resume_bytes)
