import json
import os
import logging
import io
from pathlib import Path
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from typing import Dict, List, Tuple, Optional
import re
from app.core.config import settings

logger = logging.getLogger(__name__)

class BookmarkManager:
    """Manages bookmark detection and mapping across different resume templates."""
    
    PROFILES_DIR = Path.home() / ".extract_points" / "bookmark_profiles"
    
    PATTERN_KEYWORDS = {
        'responsibilities': ['responsibilities', 'accountabilities', 'accomplishments', 'highlights'],
        'company': ['company', 'client', 'organization', 'employer'],
        'skills': ['skills', 'technical', 'technologies', 'expertise'],
        'education': ['education', 'education:', 'degree', 'certification'],
    }
    
    DEFAULT_SECTION_PATTERNS = {
        "PROFESSIONAL_EXPERIENCE": [
            "professional experience", "work experience", "experience", 
            "employment history", "career experience"
        ],
        "PROFESSIONAL_SUMMARY": [
            "professional summary", "summary", "profile", "career summary", "objective"
        ],
        "SKILLS": [
            "core competencies", "technical skills", "skills", "technologies", "key skills", "expertise"
        ],
        "EDUCATION": ["education", "academic background", "academics", "degree"],
        "CERTIFICATIONS": ["certifications", "certificates", "licenses", "training"],
    }
    
    def __init__(self):
        """Initialize bookmark manager."""
        self.PROFILES_DIR.mkdir(parents=True, exist_ok=True)
    
    def detect_bookmarks(self, resume_bytes) -> List[str]:
        """
        Auto-detect all bookmarks in a document.
        Returns list of bookmark names in order found.
        """
        try:
            # Handle path string vs BytesIO object
            if isinstance(resume_bytes, (str, Path)):
                doc = Document(resume_bytes)
            else:
                doc = Document(resume_bytes)
                
            bookmarks = []
            bookmark_count = {}
            
            for element in doc.element.iter():
                tag = element.tag
                if 'bookmarkStart' in tag:
                    if hasattr(element, 'attrib'):
                        for attr_name, attr_val in element.attrib.items():
                            if 'name' in attr_name.lower():
                                if attr_val in bookmark_count:
                                    bookmark_count[attr_val] += 1
                                    unique_name = f"{attr_val}_{bookmark_count[attr_val]}"
                                    bookmarks.append(unique_name)
                                else:
                                    bookmark_count[attr_val] = 0
                                    bookmarks.append(attr_val)
            
            return bookmarks
        except Exception as e:
            logger.error(f"Error detecting bookmarks: {e}")
            return []
    
    def match_pattern(self, bookmark_name: str) -> Tuple[str, float]:
        """Pattern match a bookmark name to determine its type."""
        bookmark_lower = bookmark_name.lower()
        best_match = ('unknown', 0.0)
        
        for pattern_type, keywords in self.PATTERN_KEYWORDS.items():
            for keyword in keywords:
                if keyword in bookmark_lower:
                    keyword_ratio = len(keyword) / len(bookmark_lower)
                    word_boundary_bonus = 0
                    if bookmark_lower.startswith(keyword):
                        word_boundary_bonus = 0.15
                    elif f'_{keyword}' in bookmark_lower or f'-{keyword}' in bookmark_lower:
                        word_boundary_bonus = 0.10
                    
                    confidence = min((keyword_ratio * 0.8) + (word_boundary_bonus), 1.0)
                    
                    if confidence > best_match[1]:
                        best_match = (pattern_type, confidence)
        
        return best_match
    
    def suggest_mappings(self, bookmarks: List[str], num_cycles: int) -> Dict[int, str]:
        """Suggest automatic cycle-to-bookmark mappings."""
        if not bookmarks:
            logger.warning("No bookmarks available for mapping suggestion")
            return {}
        
        mapping = {}
        responsibility_bookmarks = []
        other_bookmarks = []
        
        for bm in bookmarks:
            pattern_type, confidence = self.match_pattern(bm)
            if pattern_type == 'responsibilities' and confidence > 0.2:
                responsibility_bookmarks.append(bm)
            else:
                other_bookmarks.append(bm)
        
        ordered_bookmarks = responsibility_bookmarks + other_bookmarks
        
        if not ordered_bookmarks:
            ordered_bookmarks = bookmarks
        
        for cycle_num in range(1, num_cycles + 1):
            if cycle_num - 1 < len(ordered_bookmarks):
                mapping[cycle_num] = ordered_bookmarks[cycle_num - 1]
        
        logger.debug(f"Generated mapping for {len(mapping)} cycles from {len(bookmarks)} bookmarks")
        return mapping
    
    def save_profile(self, profile_name: str, bookmarks: List[str], 
                      mapping: Dict[int, str], resume_name: str = "") -> bool:
        """Save a bookmark mapping profile for future use."""
        try:
            profile_data = {
                'profile_name': profile_name,
                'resume_name': resume_name,
                'bookmarks': bookmarks,
                'mapping': {str(k): v for k, v in mapping.items()},
                'created_at': str(Path.cwd()),
            }
            
            filename = re.sub(r'[^\w\s-]', '', profile_name).strip().replace(' ', '_') + '.json'
            filepath = self.PROFILES_DIR / filename
            
            with open(filepath, 'w') as f:
                json.dump(profile_data, f, indent=2)
            
            logger.info(f"Profile saved: {profile_name}")
            return True
        except Exception as e:
            logger.error(f"Error saving profile: {e}")
            return False
    
    def load_profile(self, profile_name: str) -> Dict:
        """Load a saved bookmark profile."""
        try:
            filename = re.sub(r'[^\w\s-]', '', profile_name).strip().replace(' ', '_') + '.json'
            filepath = self.PROFILES_DIR / filename
            
            if filepath.exists():
                with open(filepath, 'r') as f:
                    data = json.load(f)
                mapping = {}
                invalid_keys = []
                for k, v in data.get('mapping', {}).items():
                    try:
                        cycle_num = int(k)
                        if cycle_num < 1:
                            invalid_keys.append(k)
                            continue
                        mapping[cycle_num] = v
                    except (ValueError, TypeError):
                        invalid_keys.append(k)
                        continue
                
                if invalid_keys:
                    logger.warning(f"Skipped {len(invalid_keys)} invalid cycle keys in profile: {invalid_keys}")
                
                data['mapping'] = mapping
                return data
        except Exception as e:
            logger.error(f"Error loading profile: {e}")
        
        return {}
    
    def list_profiles(self) -> List[Dict]:
        """List all saved bookmark profiles."""
        profiles = []
        try:
            if self.PROFILES_DIR.exists():
                for filepath in self.PROFILES_DIR.glob('*.json'):
                    try:
                        with open(filepath, 'r') as f:
                            data = json.load(f)
                            profiles.append({
                                'name': data.get('profile_name'),
                                'resume_name': data.get('resume_name', 'Unknown'),
                                'filename': filepath.name,
                                'bookmarks_count': len(data.get('bookmarks', [])),
                            })
                    except Exception as e:
                        logger.warning(f"Skipping corrupted profile {filepath.name}: {e}")
                        continue
        except Exception as e:
            logger.error(f"Error listing profiles: {e}")
        
        return profiles
    
    def delete_profile(self, profile_name: str) -> bool:
        """Delete a saved profile."""
        try:
            filename = re.sub(r'[^\w\s-]', '', profile_name).strip().replace(' ', '_') + '.json'
            filepath = self.PROFILES_DIR / filename
            
            if filepath.exists():
                filepath.unlink()
                logger.info(f"Profile deleted: {profile_name}")
                return True
        except Exception as e:
            logger.error(f"Error deleting profile: {e}")
        
        return False
    
    def find_reference_resume_path(self) -> Optional[Path]:
        """Locate a resume with existing bookmarks to use as template."""
        # 1. Check environment variables
        for env_var in ['BOOKMARK_REFERENCE_RESUME_PATH', 'BOOKMARK_TEMPLATE_PATH', 'RESUME_BOOKMARK_TEMPLATE_PATH']:
            val = os.getenv(env_var)
            if val:
                path = Path(val)
                if path.exists() and path.is_file():
                    if self.detect_bookmarks(path):
                        logger.info(f"Reference resume found via env {env_var}: {path}")
                        return path

        # 2. Check local resumes folder
        resumes_folder = settings.LOCAL_RESUMES_FOLDER
        if resumes_folder.exists():
            # First priority: files with "template" in their name inside resumes folder
            for file_path in resumes_folder.glob("**/*template*.docx"):
                if file_path.is_file() and self.detect_bookmarks(file_path):
                    logger.info(f"Template reference resume found: {file_path}")
                    return file_path
            
            # Second priority: any other *.docx files with bookmarks in resumes folder
            for file_path in resumes_folder.glob("*.docx"):
                if file_path.is_file() and self.detect_bookmarks(file_path):
                    logger.info(f"Reference resume found in resumes folder: {file_path}")
                    return file_path

        # 3. Check upload resumes folder (as fallback)
        uploads_folder = settings.UPLOAD_RESUMES_FOLDER
        if uploads_folder.exists():
            for file_path in uploads_folder.glob("**/*template*.docx"):
                if file_path.is_file() and self.detect_bookmarks(file_path):
                    logger.info(f"Template reference resume found in uploads: {file_path}")
                    return file_path
            for file_path in uploads_folder.glob("*.docx"):
                if file_path.is_file() and self.detect_bookmarks(file_path):
                    logger.info(f"Reference resume found in uploads: {file_path}")
                    return file_path

        # 4. Check workspace directories for any *.docx files with "template" in name
        workspace_dir = settings.BASE_DIR
        if workspace_dir.exists():
            for file_path in workspace_dir.glob("**/resumes/*template*.docx"):
                if file_path.is_file() and self.detect_bookmarks(file_path):
                    logger.info(f"Template resume found in base_dir: {file_path}")
                    return file_path
                    
        return None

    def _get_nonempty_paragraphs(self, doc) -> List[Tuple[int, any]]:
        """Get list of (original_idx, paragraph) for all non-empty paragraphs."""
        return [(idx, para) for idx, para in enumerate(doc.paragraphs) if para.text.strip()]

    def extract_bookmark_anchors(self, doc) -> List[Dict]:
        """Extract bookmark positions and context from reference document using non-empty paragraphs."""
        anchors = []
        nonempty_paras = self._get_nonempty_paragraphs(doc)
        
        current_heading_nonempty_idx = -1
        current_heading_text = ""
        
        for ne_idx, (orig_idx, para) in enumerate(nonempty_paras):
            if self._looks_like_heading(para):
                current_heading_text = para.text.strip()
                current_heading_nonempty_idx = ne_idx
                
            for element in para._element.iter():
                if 'bookmarkStart' in element.tag:
                    name = None
                    if hasattr(element, 'attrib'):
                        for attr_name, attr_val in element.attrib.items():
                            if 'name' in attr_name.lower():
                                name = attr_val
                                break
                    if name:
                        relative_offset = ne_idx - current_heading_nonempty_idx if current_heading_nonempty_idx != -1 else ne_idx
                        anchors.append({
                            'name': name,
                            'paragraph_index': orig_idx,
                            'paragraph_text': para.text.strip(),
                            'heading_text': current_heading_text,
                            'relative_offset': relative_offset
                        })
        logger.debug(f"Extracted {len(anchors)} bookmark anchors from reference document")
        return anchors

    def _normalize(self, text: str) -> str:
        """Standardize text for comparison."""
        if not text:
            return ""
        text = text.lower()
        text = re.sub(r'[^a-z0-9]+', ' ', text)
        return " ".join(text.split())

    def _looks_like_heading(self, paragraph) -> bool:
        """Detect if a paragraph is a section heading."""
        text = paragraph.text.strip()
        if not text:
            return False
            
        style_name = (paragraph.style.name or "").lower() if paragraph.style else ""
        if "heading" in style_name or "title" in style_name:
            return True
            
        normalized_text = self._normalize(text)
        for pattern_type, patterns in self.DEFAULT_SECTION_PATTERNS.items():
            for pat in patterns:
                if normalized_text == self._normalize(pat):
                    return True
                
        words = text.split()
        if len(words) <= 8 and text.isupper() and not text.endswith(('.', '?', '!')):
            if text.startswith(('\u2022', '-', '*', '+')) or re.match(r'^\d+\.', text):
                return False
            return True
            
        return False

    def _section_patterns_for(self, bookmark_name: str) -> List[str]:
        """Get normalized section patterns for a bookmark name."""
        name_upper = bookmark_name.upper()
        patterns = []
        for key, list_of_patterns in self.DEFAULT_SECTION_PATTERNS.items():
            if key in name_upper or name_upper in key:
                for pat in list_of_patterns:
                    patterns.append(self._normalize(pat))
        patterns.append(self._normalize(bookmark_name))
        return list(set(patterns))

    def find_matching_paragraph(self, anchor: Dict, target_paragraphs, used_paragraphs: set) -> Tuple[any, int, str, int]:
        """Find best equivalent paragraph in target document for a given bookmark using non-empty paragraph offsets."""
        best_para = None
        best_score = 0
        best_reason = "no-match"
        best_idx = -1
        
        anchor_name = anchor['name']
        anchor_heading = anchor['heading_text']
        anchor_text = anchor['paragraph_text']
        
        normalized_anchor_heading = self._normalize(anchor_heading)
        normalized_anchor_text = self._normalize(anchor_text)
        
        company_name = anchor_name.replace("_Responsibilities", "").replace("_", " ").lower()
        
        nonempty_paras = [(idx, para) for idx, para in enumerate(target_paragraphs) if para.text.strip()]
        best_ne_idx = -1
        
        for ne_idx, (idx, para) in enumerate(nonempty_paras):
            if idx in used_paragraphs:
                continue
                
            score = 0
            reasons = []
            
            para_text = para.text.strip()
            normalized_para_text = self._normalize(para_text)
            
            # 1. Section name match (+85)
            section_patterns = self._section_patterns_for(anchor_name)
            if normalized_para_text in section_patterns:
                score += 85
                reasons.append("section-name")
                
            # 2. Reference heading match (+70)
            heading_match = False
            if normalized_anchor_heading:
                if normalized_para_text == normalized_anchor_heading:
                    heading_match = True
                elif company_name in normalized_para_text and company_name in normalized_anchor_heading:
                    heading_match = True
                    
            if heading_match:
                score += 70
                reasons.append("reference-heading")
                
            # 3. Reference text match (+65)
            if normalized_anchor_text and normalized_para_text == normalized_anchor_text:
                score += 65
                reasons.append("reference-text")
                
            # 4. Bookmark terms match (+15 each, max 45)
            bookmark_terms = [t for t in self._normalize(anchor_name).split('_') if len(t) > 2]
            if not bookmark_terms:
                bookmark_terms = [t for t in self._normalize(anchor_name).split() if len(t) > 2]
            term_matches = 0
            for term in bookmark_terms:
                if term in normalized_para_text:
                    term_matches += 1
            if term_matches > 0:
                term_score = min(term_matches * 15, 45)
                score += term_score
                reasons.append(f"bookmark-terms({term_score})")
                
            # 5. Is heading (+10)
            if self._looks_like_heading(para):
                score += 10
                reasons.append("is-heading")
                
            if score > best_score:
                best_score = score
                best_para = para
                best_ne_idx = ne_idx
                best_reason = "+".join(reasons)
                best_idx = idx
                
        if best_score >= 70 and anchor.get('relative_offset', 0) > 0 and best_ne_idx != -1:
            offset_ne_idx = best_ne_idx + anchor['relative_offset']
            if offset_ne_idx < len(nonempty_paras):
                target_orig_idx, target_para = nonempty_paras[offset_ne_idx]
                if target_orig_idx not in used_paragraphs:
                    best_para = target_para
                    best_idx = target_orig_idx
                    best_reason += f"+offset({anchor['relative_offset']})"
                else:
                    logger.debug(f"Offset target index {target_orig_idx} already used.")
            else:
                logger.debug(f"Offset index {offset_ne_idx} out of range.")
                
        if best_score < 45:
            return None, 0, "no-match", -1
            
        return best_para, best_score, best_reason, best_idx

    def add_bookmark_to_paragraph(self, paragraph, bookmark_name: str, bookmark_id: int):
        """Insert Word bookmark XML elements into a paragraph safely (after w:pPr)."""
        p = paragraph._element
        
        bookmark_start = OxmlElement('w:bookmarkStart')
        bookmark_start.set(qn('w:id'), str(bookmark_id))
        bookmark_start.set(qn('w:name'), bookmark_name)
        
        bookmark_end = OxmlElement('w:bookmarkEnd')
        bookmark_end.set(qn('w:id'), str(bookmark_id))
        
        pPr = p.find(qn('w:pPr'))
        if pPr is not None:
            pPr_index = list(p).index(pPr)
            p.insert(pPr_index + 1, bookmark_start)
        else:
            p.insert(0, bookmark_start)
            
        p.append(bookmark_end)

    def ensure_bookmarks_from_reference(self, resume_bytes) -> Tuple[any, List[str], Dict]:
        """
        Automatically create bookmarks in Word documents (DOCX) that lack them
        by analyzing a reference document with existing bookmarks.
        """
        try:
            existing_bookmarks = self.detect_bookmarks(resume_bytes)
            if existing_bookmarks:
                logger.info(f"Document already has {len(existing_bookmarks)} bookmarks. No action needed.")
                if hasattr(resume_bytes, "seek"):
                    resume_bytes.seek(0)
                return resume_bytes, existing_bookmarks, {
                    "auto_created": False,
                    "created_count": 0,
                    "reference_path": None,
                    "message": "Document already has bookmarks"
                }
                
            ref_path = self.find_reference_resume_path()
            if not ref_path:
                logger.warning("No bookmarked reference resume found")
                if hasattr(resume_bytes, "seek"):
                    resume_bytes.seek(0)
                return resume_bytes, [], {
                    "auto_created": False,
                    "created_count": 0,
                    "reference_path": None,
                    "message": "No bookmarked reference resume found"
                }
                
            try:
                ref_doc = Document(ref_path)
            except Exception as e:
                logger.error(f"Could not read reference resume: {e}")
                if hasattr(resume_bytes, "seek"):
                    resume_bytes.seek(0)
                return resume_bytes, [], {
                    "auto_created": False,
                    "created_count": 0,
                    "reference_path": str(ref_path),
                    "message": f"Could not read reference resume: {str(e)}"
                }
                
            anchors = self.extract_bookmark_anchors(ref_doc)
            if not anchors:
                logger.warning("Reference resume has no usable bookmarks")
                if hasattr(resume_bytes, "seek"):
                    resume_bytes.seek(0)
                return resume_bytes, [], {
                    "auto_created": False,
                    "created_count": 0,
                    "reference_path": str(ref_path),
                    "message": "Reference resume has no usable bookmarks"
                }
                
            if hasattr(resume_bytes, "seek"):
                resume_bytes.seek(0)
                
            if isinstance(resume_bytes, (str, Path)):
                target_doc = Document(resume_bytes)
            elif isinstance(resume_bytes, bytes):
                target_doc = Document(io.BytesIO(resume_bytes))
            else:
                target_doc = Document(resume_bytes)
                
            created_bookmarks = []
            matches_metadata = []
            used_paragraphs = set()
            bookmark_id = 0
            
            target_paragraphs = target_doc.paragraphs
            
            for anchor in anchors:
                best_para, score, reason, best_idx = self.find_matching_paragraph(
                    anchor, target_paragraphs, used_paragraphs
                )
                
                if best_para:
                    self.add_bookmark_to_paragraph(best_para, anchor['name'], bookmark_id)
                    bookmark_id += 1
                    created_bookmarks.append(anchor['name'])
                    used_paragraphs.add(best_idx)
                    
                    matches_metadata.append({
                        "bookmark": anchor['name'],
                        "created": True,
                        "score": score,
                        "reason": reason
                    })
                else:
                    matches_metadata.append({
                        "bookmark": anchor['name'],
                        "created": False,
                        "score": score,
                        "reason": reason
                    })
                    
            if not created_bookmarks:
                logger.warning("No matching paragraphs found for any bookmarks")
                if hasattr(resume_bytes, "seek"):
                    resume_bytes.seek(0)
                return resume_bytes, [], {
                    "auto_created": False,
                    "created_count": 0,
                    "reference_path": str(ref_path),
                    "message": "No matching paragraphs found in target resume for any bookmarks"
                }
                
            output = io.BytesIO()
            target_doc.save(output)
            output.seek(0)
            
            logger.info(f"Successfully auto-created {len(created_bookmarks)} bookmarks using reference: {ref_path.name}")
            return output, created_bookmarks, {
                "auto_created": True,
                "created_count": len(created_bookmarks),
                "reference_path": str(ref_path),
                "matches": matches_metadata
            }
            
        except Exception as e:
            logger.error(f"Error ensuring bookmarks: {e}", exc_info=True)
            if hasattr(resume_bytes, "seek"):
                resume_bytes.seek(0)
            return resume_bytes, [], {
                "auto_created": False,
                "created_count": 0,
                "reference_path": None,
                "message": f"Error ensuring bookmarks: {str(e)}"
            }

    def validate_mapping(self, mapping: Dict[int, str], 
                        available_bookmarks: List[str]) -> Tuple[bool, str]:
        """Validate that a mapping uses only available bookmarks."""
        for cycle_num, bookmark in mapping.items():
            if bookmark not in available_bookmarks:
                return False, f"Bookmark '{bookmark}' not found in document"
        
        return True, ""
